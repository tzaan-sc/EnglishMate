"""
document_parser.py
===================
Trích xuất tự động Câu hỏi Trắc nghiệm và Từ vựng từ file Word (.docx), PDF (.pdf), Text (.txt, .md).
"""

import os
import re
import zipfile
import xml.etree.ElementTree as ET


def extract_text_from_file(file_path: str) -> str:
    """Đọc và trích xuất toàn bộ văn bản thuần từ file Word, PDF hoặc TXT."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in [".txt", ".md", ".csv", ".json"]:
        for enc in ["utf-8-sig", "utf-8", "cp1258", "latin-1"]:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except Exception:
                continue
        return ""

    elif ext == ".docx":
        # 1. Thử dùng thư viện python-docx
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        full_text.append(row_text)
            return "\n".join(full_text)
        except Exception:
            pass

        # 2. Fallback dùng zipfile + XML nếu không có docx
        try:
            with zipfile.ZipFile(file_path) as z:
                xml_content = z.read("word/document.xml")
            tree = ET.fromstring(xml_content)
            namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            paragraphs = []
            for p in tree.findall(".//w:p", namespaces):
                texts = [node.text for node in p.findall(".//w:t", namespaces) if node.text]
                if texts:
                    paragraphs.append("".join(texts))
            return "\n".join(paragraphs)
        except Exception as e:
            print(f"[Lỗi đọc docx] {file_path}: {e}")
            return ""

    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                txt = page.extract_text()
                if txt:
                    pages_text.append(txt)
            return "\n".join(pages_text)
        except Exception as e:
            print(f"[Lỗi đọc pdf] {file_path}: {e}")
            return ""

    return ""


# ====================================================================
# PHẦN 1: BÓC TÁCH CÂU HỎI TRẮC NGHIỆM (MULTIPLE CHOICE QUESTIONS)
# ====================================================================

def parse_questions_from_text(raw_text: str, default_topic="General", default_level="B1", default_skill="Grammar") -> list:
    """
    Nhận diện thông minh các khối câu hỏi trắc nghiệm tiếng Anh.
    Hỗ trợ các dạng đánh số:
      Câu 1:, Question 1., 1., 1/, [1], Part 5: Question 101...
    Hỗ trợ đáp án A, B, C, D trên cùng dòng hoặc xuống dòng.
    Hỗ trợ bóc tách Đáp án đúng (Key/Answer) và Giải thích (Explanation).
    """
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    full_clean_text = "\n".join(lines)

    # 1. Tách danh sách đáp án tổng hợp (nếu có bảng key ở cuối tài liệu, ví dụ: 1.A 2.B 3.C 4.D)
    answer_key_map = {}
    key_section_match = re.search(r"(?:ĐÁP ÁN|DAP AN|ANSWER KEY|KEY|BẢNG ĐÁP ÁN)[:\s\n]+([\s\S]+)$", full_clean_text, re.IGNORECASE)
    if key_section_match:
        key_text = key_section_match.group(1)
        for m in re.finditer(r"(?:Câu|Question|\b)?\s*(\d+)[\.\s\:\-\)]+\s*([A-D])\b", key_text, re.IGNORECASE):
            answer_key_map[int(m.group(1))] = m.group(2).upper()

    # 2. Tách các câu hỏi bằng regex chia khối
    # Regex nhận diện điểm bắt đầu câu hỏi
    q_split_pattern = r"(?=(?:^|\n)\s*(?:Câu|Question|\b)\s*(\d+)[\.\:\/\-\)])"
    chunks = re.split(q_split_pattern, full_clean_text, flags=re.IGNORECASE)

    questions = []
    
    # chunks sẽ có dạng: [preamble, q_num_1, chunk_1, q_num_2, chunk_2, ...]
    if len(chunks) > 1 and chunks[1].isdigit():
        idx = 1
        while idx < len(chunks) - 1:
            q_num = int(chunks[idx])
            q_body = chunks[idx + 1]
            idx += 2
            
            parsed_q = _extract_single_question(q_num, q_body, answer_key_map, default_topic, default_level, default_skill)
            if parsed_q:
                questions.append(parsed_q)
    else:
        # Cách 2: Quét tuần tự từng dòng nếu không chia chunk được
        current_q_num = None
        current_lines = []
        for line in lines:
            m = re.match(r"^(?:Câu|Question|\b)?\s*(\d+)[\.\:\/\-\)]\s*(.*)", line, re.IGNORECASE)
            if m and len(m.group(1)) <= 4:
                if current_q_num is not None and current_lines:
                    parsed_q = _extract_single_question(current_q_num, "\n".join(current_lines), answer_key_map, default_topic, default_level, default_skill)
                    if parsed_q:
                        questions.append(parsed_q)
                current_q_num = int(m.group(1))
                current_lines = [m.group(2) if m.group(2) else line]
            else:
                if current_q_num is not None:
                    current_lines.append(line)
        
        if current_q_num is not None and current_lines:
            parsed_q = _extract_single_question(current_q_num, "\n".join(current_lines), answer_key_map, default_topic, default_level, default_skill)
            if parsed_q:
                questions.append(parsed_q)

    return questions


def _extract_single_question(q_num: int, text: str, key_map: dict, topic: str, level: str, skill: str) -> dict:
    """Bóc tách chi tiết Question Text, Option A, B, C, D, Correct Option, Explanation."""
    lines = text.strip().splitlines()
    if not lines:
        return None

    # Tìm đáp án đúng trực tiếp trong câu (nếu có: Đáp án: A, Key: B, Chọn C...)
    correct_option = key_map.get(q_num, "")
    explanation = ""

    ans_match = re.search(r"(?:Đáp án|Dap an|Answer|Key|Chọn|Đ/a)[\s\:\-\.]*([A-D])\b", text, re.IGNORECASE)
    if ans_match:
        correct_option = ans_match.group(1).upper()

    exp_match = re.search(r"(?:Giải thích|Giai thich|Explanation|Lý do|Note)[\s\:\-\.]([\s\S]+)", text, re.IGNORECASE)
    if exp_match:
        explanation = exp_match.group(1).strip()
        # Loại bỏ phần explanation ra khỏi text để đỡ nhiễu options
        text = text[:exp_match.start()].strip()

    # Tìm Options A, B, C, D
    # Pattern 1: Tìm A. ... B. ... C. ... D. ... trên cùng 1 dòng hoặc các dòng
    opt_pattern = r"(?:^|\s|\t)(?:[\[\(]?([A-D])[\]\)\.\:\-])\s*([\s\S]*?)(?=(?:[\[\(]?[A-D][\]\)\.\:\-]|$))"
    options_found = {}

    for m in re.finditer(opt_pattern, text):
        opt_letter = m.group(1).upper()
        opt_val = m.group(2).strip()
        # Clean trailing answer tag if present
        opt_val = re.sub(r"(?:Đáp án|Dap an|Answer|Key|Giải thích)[\s\S]*$", "", opt_val, flags=re.IGNORECASE).strip()
        if opt_val and opt_letter not in options_found:
            options_found[opt_letter] = opt_val

    # Question text là phần trước Option đầu tiên
    first_opt_match = re.search(r"(?:^|\s)(?:[\[\(]?[A-D][\]\)\.\:\-])\s*", text)
    if first_opt_match:
        q_text = text[:first_opt_match.start()].strip()
    else:
        q_text = lines[0].strip()

    # Clean q_text: Xóa tiền tố "Câu 1:", "Question 1."
    q_text = re.sub(r"^(?:Câu|Question|\b)?\s*\d+[\.\:\/\-\)]\s*", "", q_text, flags=re.IGNORECASE).strip()

    if not q_text or len(options_found) < 2:
        return None

    # Mặc định nếu thiếu A, B, C, D thì gán rỗng
    opt_a = options_found.get("A", "")
    opt_b = options_found.get("B", "")
    opt_c = options_found.get("C", "")
    opt_d = options_found.get("D", "")

    if not correct_option:
        correct_option = "A"  # Mặc định an toàn nếu chưa gán key

    if not explanation:
        explanation = f"Đáp án chính xác là {correct_option}."

    return {
        "question_number": q_num,
        "question_text": q_text,
        "option_a": opt_a,
        "option_b": opt_b,
        "option_c": opt_c,
        "option_d": opt_d,
        "correct_option": correct_option,
        "explanation": explanation,
        "topic": topic,
        "level": level,
        "skill": skill
    }


# ====================================================================
# PHẦN 2: BÓC TÁCH TỪ VỰNG (VOCABULARY LIST PARSER)
# ====================================================================

def parse_vocabulary_from_text(raw_text: str, default_topic="General", default_level="A2") -> list:
    """
    Bóc tách từ vựng từ văn bản.
    Nhận diện các mẫu phổ biến:
      - accomplish /əˈkʌm.plɪʃ/ (v): hoàn thành, đạt được - She accomplished her goal.
      - resilient (adj) - kiên cường
      - negotiate (v): đàm phán
      - 1. innovate (v): đổi mới
    """
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    vocab_list = []

    # Regex nhận diện dòng từ vựng
    # Group 1: Word
    # Group 2: Pronunciation (tùy chọn trong /.../)
    # Group 3: Part of speech (tùy chọn trong (...))
    # Group 4: Meaning & example
    vocab_pattern = r"^(?:\d+[\.\-\)]\s*|\-\s*|\*\s*)?([a-zA-Z\s\-\'\’]+?)(?:\s+[\/\[]([^\/\]]+)[\/\]])?(?:\s*\(([a-zA-Z\.\,\s]+)\))?\s*(?:[\:\-\=]\s*|\t)(.+)$"

    for idx, line in enumerate(lines, start=1):
        m = re.match(vocab_pattern, line)
        if m:
            word = m.group(1).strip()
            pronunciation = m.group(2).strip() if m.group(2) else ""
            pos_raw = m.group(3).strip().lower() if m.group(3) else "noun"
            rest = m.group(4).strip()

            # Chuẩn hóa Part of speech
            pos = "noun"
            if any(k in pos_raw for k in ["v", "verb", "động từ"]):
                pos = "verb"
            elif any(k in pos_raw for k in ["adj", "a.", "tính từ"]):
                pos = "adjective"
            elif any(k in pos_raw for k in ["adv", "phó từ", "trạng từ"]):
                pos = "adverb"
            elif any(k in pos_raw for k in ["n", "noun", "danh từ"]):
                pos = "noun"

            # Tách meaning và example nếu có dấu gạch ngang hoặc ví dụ (Ex:)
            meaning = rest
            example_en = f"This is an example using the word '{word}'."
            example_vi = f"Đây là ví dụ minh họa cho từ '{word}'."

            ex_match = re.search(r"(?:Ex|E\.g|Ví dụ|Vi du)[\s\:\-\.]([^\—\-\n]+)(?:[\—\-\:]\s*(.+))?", rest, re.IGNORECASE)
            if ex_match:
                meaning = rest[:ex_match.start()].strip(" -:;,")
                example_en = ex_match.group(1).strip()
                if ex_match.group(2):
                    example_vi = ex_match.group(2).strip()
                else:
                    example_vi = f"Ví dụ cho câu: {example_en}"

            if word and len(word) <= 60 and not word.startswith("http"):
                vocab_list.append({
                    "word": word,
                    "pronunciation": f"/{pronunciation}/" if pronunciation and not pronunciation.startswith("/") else (pronunciation or f"/{word}/"),
                    "part_of_speech": pos,
                    "meaning_vi": meaning or "Nghĩa đang cập nhật",
                    "example_en": example_en,
                    "example_vi": example_vi,
                    "topic": default_topic,
                    "level": default_level,
                    "image_url": "",
                    "collocations": "",
                    "synonyms": "",
                    "antonyms": ""
                })

    return vocab_list
